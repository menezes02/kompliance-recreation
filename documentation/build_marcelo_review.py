from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Kompliance_Marcelo_Requirements_Delivery_Status_2026-07-22.docx"

NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "0F766E"
GREEN = "DCFCE7"
GREEN_TEXT = "166534"
BLUE_FILL = "DBEAFE"
BLUE_TEXT = "1E40AF"
AMBER = "FEF3C7"
AMBER_TEXT = "92400E"
RED = "FEE2E2"
RED_TEXT = "991B1B"
GREY = "F2F4F7"
MID_GREY = "D0D5DD"
DARK = "27364A"
WHITE = "FFFFFF"
BODY = "344054"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width: int = 9360, indent: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GREY, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, fld_char_end])


def add_bottom_rule(paragraph, color=TEAL, size="20") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


STATUS_STYLE = {
    "Implemented": (GREEN, GREEN_TEXT),
    "Substantially complete": (BLUE_FILL, BLUE_TEXT),
    "Partially complete": (AMBER, AMBER_TEXT),
    "Deferred for private pilot": (GREY, DARK),
    "Pending customer approval": (AMBER, AMBER_TEXT),
    "External dependency": (RED, RED_TEXT),
}


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string("1F4D78")
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(BODY)
        style.paragraph_format.space_after = Pt(3)

    props = doc.core_properties
    props.title = "Kompliance – Marcelo Requirements & Delivery Status"
    props.subject = "Traceable response to Marcelo's platform review"
    props.author = "Kompliance Project Team"
    props.keywords = "Kompliance, requirements, delivery, review, pilot, compliance"
    props.comments = "Prepared from the original Marcelo review and verified project records."


def add_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("KOMPLIANCE  |  MARCELO DELIVERY REVIEW")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_bottom_rule(p, MID_GREY, "6")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential — controlled customer pilot  •  Page ")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("667085")
    add_page_field(p, "PAGE")
    p.add_run(" of ")
    add_page_field(p, "NUMPAGES")


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("KOMPLIANCE PLATFORM")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    r.font.all_caps = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Marcelo Review")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Requirements & Delivery Status")
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    add_bottom_rule(p, TEAL, "20")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2100, 7260]
    metadata = [
        ("Prepared for", "Marcelo — requirements and customer review"),
        ("Status date", "22 July 2026"),
        ("Live environment", "kompliance.felipeitprojects.com"),
        ("Deployed release", "d96156a7c3b — controlled pilot"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, top=55, bottom=55)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], GREY)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        rr = p0.add_run(label.upper())
        rr.font.size = Pt(8)
        rr.font.bold = True
        rr.font.color.rgb = RGBColor.from_string("667085")
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        rr = p1.add_run(value)
        rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor.from_string(DARK)
    set_table_borders(table, "E4E7EC", "2")


def add_callout(doc: Document, heading: str, text: str, fill=BLUE_FILL, accent=BLUE_TEXT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_width(cell, 9360)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(heading.upper())
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(accent)
    p = cell.add_paragraph(text)
    p.paragraph_format.space_after = Pt(0)
    r = p.runs[0]
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    set_table_borders(table, fill, "2")


def add_status_summary(doc: Document) -> None:
    doc.add_heading("Executive status", level=1)
    add_callout(
        doc,
        "Review conclusion",
        "The single-company private pilot is ready for Marcelo and the customer to test. Core engineering is substantially complete; commercial release still requires acceptance, external verification and governance approvals.",
    )
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2300, 1560, 5500]
    headers = ["Measure", "Position", "Meaning"]
    for i, (cell, width, label) in enumerate(zip(table.rows[0].cells, widths, headers)):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Customer pilot", "≈95%", "Deployed and ready for controlled acceptance testing."),
        ("Full Marcelo scope", "≈80%", "Most product foundations are implemented; public-demo, provider and scale work remains."),
        ("Commercial release", "Pending", "Requires named sign-off, privacy/legal approval, Google verification and independent security assurance."),
    ]
    for measure, position, meaning in rows:
        row = table.add_row()
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[0].text = measure
        row.cells[1].text = position
        row.cells[2].text = meaning
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(TEAL)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9)
    set_table_borders(table)

    p = doc.add_paragraph()
    r = p.add_run("Estimate basis. ")
    r.bold = True
    p.add_run("Percentages are indicative delivery estimates from the current gap analysis, not contractual completion certificates. Customer testing can change them.")


def add_status_legend(doc: Document) -> None:
    doc.add_heading("Status legend", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2600, 3380, 3380]
    entries = [
        ("Implemented", "Delivered and evidenced in the current release."),
        ("Substantially complete", "Core path delivered; limited validation/refinement remains."),
        ("Partially complete", "Useful capability exists, but part of Marcelo's request remains."),
        ("Deferred for private pilot", "Intentionally held because the pilot uses authorised customer data."),
        ("Pending customer approval", "Built or prepared, but needs Marcelo/customer decision or sign-off."),
        ("External dependency", "Depends on a provider, legal process or independent assessor."),
    ]
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, GREY)
        set_cell_margins(cell)
        cell.paragraphs[0].text = "Status" if i == 0 else ("Interpretation" if i == 1 else "Status / interpretation")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for idx in range(0, len(entries), 2):
        row = table.add_row()
        prevent_row_split(row)
        left = entries[idx]
        right = entries[idx + 1]
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
        status, meaning = left
        row.cells[0].text = status
        fill, color = STATUS_STYLE[status]
        set_cell_shading(row.cells[0], fill)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)
        row.cells[1].text = meaning
        status2, meaning2 = right
        row.cells[2].text = f"{status2}\n{meaning2}"
        fill2, color2 = STATUS_STYLE[status2]
        set_cell_shading(row.cells[2], fill2)
        row.cells[2].paragraphs[0].runs[0].font.bold = True
        row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color2)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(8.5)
    set_table_borders(table)


MODULES = [
    {
        "title": "Module 1 — Universal Worker Profile",
        "objective": "Create a universal digital worker profile shareable across companies and induction systems.",
        "estimate": "94%",
        "rows": [
            ("1.1", "Worker registration", "Free self-registration, email verification, password recovery, secure sessions, lockout, TOTP MFA and one-time backup codes are implemented.", "Partially complete", "Mobile/SMS authentication needs an approved provider and privacy terms."),
            ("1.2", "Worker profile", "Editable profile supports identity, contact, emergency, employment, skills, trade, qualifications, optional medical data, certifications, training and uploaded documents.", "Substantially complete", "Marcelo/customer must validate the final field set and terminology."),
            ("1.3", "QR code", "Each worker receives a unique QR/public profile. Companies can scan with camera or paste the link, request access and import approved fields.", "Implemented", "Complete the customer QR journey during acceptance testing."),
            ("1.4", "Profile sharing", "Field-level consent, QR, secure links, API sharing, company relationships and immediate worker revocation are implemented and audited.", "Implemented", "Confirm consent wording with the privacy owner."),
            ("1.5", "API integration", "Audited REST resources cover profiles, certifications, documents, inductions and training. OpenAPI 3.1, bearer-token revocation, pagination and rate limits are present.", "Substantially complete", "A client SDK and partner-specific integration test follow an approved integration."),
            ("1.6", "Central database", "Workers, certifications, training, documents and company relationships are centralised with tenant scoping and protected-source boundaries.", "Implemented", "Independent multi-tenant security testing remains a commercial-release gate."),
            ("1.7", "Multi-language support", "Persistent English, Portuguese and Spanish support covers the shell, worker portal and primary workflows; the architecture supports expansion.", "Partially complete", "Native-speaker review and rare/error-state strings remain."),
        ],
    },
    {
        "title": "Module 2 — Document Management",
        "objective": "Allow workers and companies to upload, manage, review and track compliance documents.",
        "estimate": "84%",
        "rows": [
            ("2.1", "Upload system", "Drag-and-drop, browse, multiple-file upload, progress, status, validation and accepted-format controls are implemented for writable workflows.", "Implemented", "Imported source documents remain deliberately read-only."),
            ("2.2", "Document storage", "GA1/GA2/GA3 and broader worker/company document libraries are available; worker-owned categories and tenant review are implemented.", "Partially complete", "Plant, equipment and site ownership refinements remain for the full product model."),
            ("2.3", "Document viewer", "PDF/image preview, browser view, download, replace/delete for local writable records and automatic version history are implemented.", "Substantially complete", "Improve visual version grouping; protected imported files cannot be replaced or deleted."),
            ("2.4", "Expiry management", "Expiry dates, valid/due-soon/expired colour states, expiry centre, reminder preparation and auditable best-effort extraction are implemented.", "Partially complete", "Scanned-image OCR/provider extraction and approved automatic scheduling remain."),
        ],
    },
    {
        "title": "Module 3 — Compliance, Certifications & Workflow",
        "objective": "Automate compliance reviews, approvals, requests and communication.",
        "estimate": "82%",
        "rows": [
            ("3.1", "Internal messaging", "Worker/company conversations and routed requests support GA forms, inductions, certifications and operational topics.", "Substantially complete", "Broader attachment/link types are a refinement."),
            ("3.2", "Requests", "Inspection, renewal, approval, missing-document and information requests can be routed and tracked.", "Implemented", "Customer to validate labels and priorities."),
            ("3.3", "Department contacts", "Tenant company/contact settings and responsible department routing are implemented.", "Implemented", "Marcelo/customer must enter and approve production contacts."),
            ("3.4", "Notifications", "In-app notifications, preferences, queued email, retry/history and Gmail API send capability are implemented. A controlled live email was accepted by Gmail.", "Partially complete", "SMS/push providers are not selected; automated scheduler is intentionally disabled pending approval."),
            ("3.5", "Certification & training status", "Valid, due-soon and expired states are shown across the expiry centre, certificates, training and relevant documents.", "Implemented", "Customer acceptance must confirm business rules and reminder windows."),
            ("3.6", "Supervisor review status", "Unread/viewed and approve/decline history records who uploaded, first viewed and reviewed each tenant document.", "Implemented", "Validate the supervisor queue with real pilot roles."),
            ("3.7", "Induction approval workflow", "Pending/approved/declined decisions, comments, additional-information requests, reviewer/date and immutable history are implemented.", "Implemented", "Complete the end-to-end acceptance scenario."),
            ("3.8", "Worker communication", "In-app and email pathways cover approval, decline, requests, assignments and expiry events with user preferences.", "Partially complete", "SMS remains an external dependency; production reminder recipients and intervals need approval."),
        ],
    },
    {
        "title": "Module 4 — Responsive User Interface",
        "objective": "Provide usable desktop, tablet and mobile experiences for dashboards, tables, forms and document workflows.",
        "estimate": "95%",
        "rows": [
            ("4.1", "Responsive components", "Navigation, dashboards, tables, employee lists, forms, dialogs and document viewers adapt to desktop and narrow mobile widths.", "Substantially complete", "Exhaustive tablet/device and touch acceptance remains."),
            ("4.2", "Interaction quality", "Spacing, resizing, scrolling, filters, one-calendar date ranges and mobile-friendly controls are implemented in primary paths.", "Substantially complete", "Marcelo/customer should report clipping, scrolling, labels and touch issues during UAT."),
            ("4.3", "Accessibility/localisation quality", "Primary journeys were checked at desktop and 390 px without browser-console errors; language preference persists.", "Partially complete", "Formal accessibility audit and native-language review remain."),
        ],
    },
    {
        "title": "Module 5 — Demo Environment",
        "objective": "Create a safe sales/demo environment containing only fictional information and watermarked documents.",
        "estimate": "5%",
        "rows": [
            ("5.1", "Fictional tenant and data", "Governance boundaries are documented, but the current instruction was to use the authorised customer snapshot for a private pilot.", "Deferred for private pilot", "Build a completely separate fictional tenant before any public or sales demonstration."),
            ("5.2", "Document redaction and watermarking", "No demo-document conversion was performed because the pilot must preserve authorised source records unchanged.", "Deferred for private pilot", "Create synthetic files, apply SAMPLE/DEMO watermarks and run a PII scan before public use."),
            ("5.3", "Public-demo approval", "The present environment is labelled and controlled as a private customer pilot, not a public demo.", "Pending customer approval", "Marcelo must decide whether a separate public demo is required for the next phase."),
        ],
    },
    {
        "title": "Module 6 — Data Migration",
        "objective": "Extract authorised client data and import it while preserving all worker/company relationships.",
        "estimate": "82%",
        "rows": [
            ("6.1", "Current authorised snapshot", "One customer snapshot is imported as 3,597 protected records, with 3,077 source files held read-only and isolated from writable pilot records.", "Implemented", "Marcelo/customer should reconcile representative totals and documents during UAT."),
            ("6.2", "Repeatable migration tooling", "Checksum inventories, dry-run reconciliation, transactional tenant import, relationship remapping, provenance, replay prevention and run history are implemented.", "Substantially complete", "Produce a customer-specific acceptance report for each migration."),
            ("6.3", "Additional clients", "The platform can accept isolated tenant packages, but no other client may be extracted without written authorisation and source mapping.", "Pending customer approval", "Approve each client, schema, lawful basis, transfer method and reconciliation plan."),
        ],
    },
    {
        "title": "Module 7 — Business Objectives",
        "objective": "Reduce repeated entry and onboarding time while improving compliance visibility and cross-company portability.",
        "estimate": "55%",
        "rows": [
            ("7.1", "Reduced repeat entry", "Central worker profiles, QR identity, consented company access and profile import are implemented.", "Substantially complete", "Measure actual onboarding time and repeat-entry reduction in pilot use."),
            ("7.2", "Improved compliance management", "Central dashboards, expiry states, document review, audit history and supervisor workflows are implemented.", "Substantially complete", "Customer acceptance and operational ownership remain."),
            ("7.3", "Shared worker network", "Multi-company consent, revocation and API foundations exist.", "Partially complete", "Network adoption, onboarding/support processes and tenant assurance are not yet proven."),
            ("7.4", "International expansion", "Three-language architecture and primary catalogue are present.", "Partially complete", "Native review, additional locales and international privacy/operational readiness remain."),
        ],
    },
]


def add_requirement_table(doc: Document, module: dict) -> None:
    heading = doc.add_heading(module["title"], level=1)
    keep_with_next(heading)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run("Original objective. ")
    r.bold = True
    p.add_run(module["objective"])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Indicative module position: {module['estimate']}")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [600, 1500, 3260, 1700, 2300]
    headers = ["Ref.", "Requirement", "What was delivered", "Status", "Remaining / review action"]
    for cell, width, label in zip(table.rows[0].cells, widths, headers):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])

    for ref, requirement, delivered, status, remaining in module["rows"]:
        row = table.add_row()
        prevent_row_split(row)
        values = [ref, requirement, delivered, status, remaining]
        for cell, width, value in zip(row.cells, widths, values):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = value
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string(BODY)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        fill, color = STATUS_STYLE[status]
        set_cell_shading(row.cells[3], fill)
        row.cells[3].paragraphs[0].runs[0].font.bold = True
        row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)
    set_table_borders(table)


def add_completed_work(doc: Document) -> None:
    doc.add_heading("What has been completed", level=1)
    doc.add_paragraph("The following release-candidate work is present in the current project and underpins the requirement statuses above:")
    groups = [
        ("Data protection and tenancy", [
            "Protected customer snapshot separated from writable pilot records; imported source rows cannot be edited or deleted.",
            "Company tenancy, role-scoped accounts, consent, revocation, audit history and protected archive boundaries.",
            "Repeatable migration packages with checksums, dry-run reconciliation, provenance and replay prevention.",
        ]),
        ("Worker and company workflows", [
            "Self-registration, verification, recovery, MFA, editable worker profile, QR/public profile and field-level sharing.",
            "Company QR access request, worker approval/decline, workforce import/refresh and immediate revocation.",
            "Document upload/view/version/expiry, supervisor review, routed requests, conversations and induction decisions.",
        ]),
        ("Forms, certificates and compliance", [
            "Assignments, drafts, required-field validation, signatures, evidence, submissions and generated PDF reports.",
            "Numbered certificate PDFs, QR verification, replacement, revocation and expiry status.",
            "Expiry centre, configurable reminder windows, deduplication, notification history and retries.",
        ]),
        ("Security and operations", [
            "Secure sessions, CSRF protection, account lockout, one-use recovery links, session revocation and audit logging.",
            "Read-only container filesystems, non-root application user, no-new-privileges, health checks and operations monitoring.",
            "Checksum-verified backups, restore rehearsal, documented rollback points and automatic writable-data backups.",
        ]),
        ("Production communication", [
            "Gmail API OAuth configured with the minimum gmail.send scope; a controlled message from the live container was accepted.",
            "OAuth application is in production mode, while formal Google verification remains outstanding.",
            "Automatic reminder scheduling remains deliberately disabled until recipient, interval and governance approval.",
        ]),
    ]
    for title, bullets in groups:
        doc.add_heading(title, level=2)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")


def add_deployment_evidence(doc: Document) -> None:
    doc.add_heading("Current deployment evidence", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2600, 2360, 4400]
    headers = ["Control", "Verified position", "Evidence / significance"]
    for cell, width, label in zip(table.rows[0].cells, widths, headers):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Live release", "d96156a7c3b", "Application is deployed to the controlled HTTPS pilot environment."),
        ("Container health", "3 services healthy", "Application, gateway and operations services passed deployment checks."),
        ("Protected data", "3,597 records", "All retain the protected production-source marker; source archive is read-only."),
        ("Source archive", "3,077 files", "Pilot can view authorised files while preserving the original archive."),
        ("Email", "Controlled send accepted", "Gmail accepted a message sent from the live application container."),
        ("Scheduler", "Disabled by design", "Prevents automatic outbound reminders before customer approval."),
        ("Backups", "Pre/post verified", "Deployment and writable-data rollback paths are recorded with SHA-256 checksums."),
    ]
    for control, position, evidence in rows:
        row = table.add_row()
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, (control, position, evidence)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    set_table_borders(table)
    p = doc.add_paragraph()
    p.add_run("Security note. ").bold = True
    p.add_run("Passwords, OAuth values, tokens, SSH key material and recovery codes are intentionally excluded from this review document.")


def add_review_actions(doc: Document) -> None:
    doc.add_heading("What Marcelo should review now", level=1)
    actions = [
        "Run the primary customer path: dashboard, GA1/GA2/GA3 filters, one-date/two-date range selection, PDF view/download and protected-record checks.",
        "Run one complete local workflow: assignment, draft/resume, signature, evidence, final submission and generated PDF report.",
        "Run the worker consent path: register, complete profile, create QR link, request company access, approve selected fields and revoke access.",
        "Validate supervisor operations: document review, induction approve/decline/comment, routed request and notification history.",
        "Repeat the primary paths on desktop, tablet and mobile; record browser/device details for any issue.",
        "Approve or correct company contacts, privacy contact, retention wording, notification recipients and reminder intervals.",
        "Record each defect with page/action, expected result, actual result, screenshot with personal information removed and severity.",
    ]
    for action in actions:
        doc.add_paragraph(action, style="List Number")

    add_callout(
        doc,
        "Acceptance decision",
        "The appropriate decision after testing is Accepted, Accepted with conditions, or Rejected with blocking defects. Commercial launch should not be inferred from pilot acceptance alone.",
        fill=AMBER,
        accent=AMBER_TEXT,
    )


def add_remaining_and_plan(doc: Document) -> None:
    doc.add_heading("Remaining work before 100%", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2100, 2500, 2460, 2300]
    headers = ["Workstream", "What remains", "Dependency / owner", "Release impact"]
    for cell, width, label in zip(table.rows[0].cells, widths, headers):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    rows = [
        ("Customer acceptance", "Complete checklist, triage defects and obtain named sign-off.", "Marcelo, pilot user, technical owner", "Blocks customer-approved release"),
        ("Notifications", "Approve recipients/intervals; select SMS/push providers if required; then enable scheduler.", "Business/privacy owner", "Email manual path works; automation held"),
        ("Google verification", "Submit the Gmail sensitive-scope verification evidence and complete Google's process.", "Project owner / Google", "Blocks removal of unverified warning and 100-user cap"),
        ("Security assurance", "Independent penetration test and multi-tenant/privacy review.", "External assessor", "Blocks broad commercial rollout"),
        ("Legal and governance", "DPA, lawful basis, retention, subprocessors, incident/support ownership and Irish/EU review.", "Legal/privacy/business owners", "Blocks commercial operation"),
        ("Public demo", "Build fictional tenant, synthetic/watermarked files and automated PII checks if a public demo is wanted.", "Product owner", "Not required for the current private pilot"),
        ("Product polish", "Native-language review, accessibility/device testing, OCR and client SDK as prioritised.", "Product/engineering", "Quality and scale improvements"),
    ]
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, values):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.size = Pt(8)
    set_table_borders(table)

    doc.add_heading("Recommended path to release", level=1)
    steps = [
        ("1. Accept the controlled pilot", "Complete Marcelo/customer UAT and close blocking defects."),
        ("2. Approve governance", "Confirm owners, privacy wording, retention, providers, recipients and support process."),
        ("3. Enable approved automation", "Run a controlled reminder test, inspect delivery history, then explicitly enable the scheduler."),
        ("4. Obtain external assurance", "Complete Google verification, penetration testing and legal/privacy review."),
        ("5. Release by scope", "Approve single-company production first; approve multi-company/public-demo operation only after its additional gates."),
    ]
    for title, detail in steps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title + " — ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        p.add_run(detail)


def add_evidence_register(doc: Document) -> None:
    doc.add_heading("Evidence register", level=1)
    doc.add_paragraph("This response was prepared from Marcelo’s original review and the following project-controlled evidence. These files remain in the same main project folder as the application source and backups.")
    evidence = [
        ("MARCELO_GAP_ANALYSIS.md", "Current module estimates, implemented scope, gaps and milestone sequence."),
        ("PILOT_TEST_HANDOFF.md", "Controlled test paths, protected-data boundary and deliberate pilot holds."),
        ("PILOT_ACCEPTANCE_CHECKLIST.md", "Formal customer validation and sign-off checklist."),
        ("RELEASE_CHECKLIST.md", "Release controls, configuration, smoke tests, backup and rollback requirements."),
        ("OPERATIONS_RUNBOOK.md", "Health checks, backups, restore rehearsal, monitoring, email/scheduler and privacy operations."),
        ("APP_MAP.md", "Mapped functions, screens, routes, fields and observed application behaviour."),
        ("UNIVERSAL_WORKER_API.md", "Universal worker REST resources and integration safeguards."),
        ("SUPERVISOR_WORKFLOW.md", "Requests, messaging, review and induction approval behaviour."),
        ("TENANT_MIGRATION.md", "Checksum-inventoried migration and reconciliation process."),
        ("workstation-backups/releases/DEPLOYMENT-d96156a.md", "Deployed commit, live verification, checksums, rollback locations and Gmail status."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [3600, 5760]
    for cell, width, label in zip(table.rows[0].cells, widths, ("Evidence file", "Purpose")):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for filename, purpose in evidence:
        row = table.add_row()
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, (filename, purpose)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    set_table_borders(table)

    doc.add_heading("Review record", level=1)
    review = doc.add_table(rows=4, cols=2)
    review.alignment = WD_TABLE_ALIGNMENT.LEFT
    review.autofit = False
    set_table_width(review)
    widths = [2600, 6760]
    fields = [
        ("Current decision", "Pending Marcelo/customer review"),
        ("Decision options", "Accepted  |  Accepted with conditions  |  Rejected with blocking defects"),
        ("Conditions / defects", "Record in PILOT_ACCEPTANCE_CHECKLIST.md and the agreed issue tracker."),
        ("Required approvers", "Business owner, pilot user and technical owner."),
    ]
    for row, values in zip(review.rows, fields):
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, values):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=100, bottom=100)
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], GREY)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_table_borders(review)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("END OF REVIEW RESPONSE")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("667085")


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_running_header_footer(doc)
    add_masthead(doc)
    add_status_summary(doc)
    doc.add_page_break()
    add_status_legend(doc)
    doc.add_heading("Scope and reading guide", level=1)
    doc.add_paragraph(
        "This document recreates Marcelo’s original review as a traceable delivery response. Each original numbered requirement group is retained, then matched to the capability delivered, its present status and the action still required. Repeated feature notes at the end of the original review are consolidated into the relevant numbered modules so they are not counted twice."
    )
    p = doc.add_paragraph()
    p.add_run("Private-pilot boundary. ").bold = True
    p.add_run(
        "The current environment uses an authorised customer snapshot and is not a fictional public demo. All imported records remain protected and read-only. Local pilot workflows may create separate writable records without changing the snapshot."
    )
    doc.add_heading("Requirement-by-requirement response", level=1)
    doc.add_paragraph("The module estimates below are the current gap-analysis positions. Row-level status is the operative review detail.")
    for module in MODULES:
        add_requirement_table(doc, module)
    add_completed_work(doc)
    add_deployment_evidence(doc)
    add_review_actions(doc)
    add_remaining_and_plan(doc)
    add_evidence_register(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(path)
