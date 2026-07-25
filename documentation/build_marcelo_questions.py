from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from build_marcelo_review import (
    AMBER,
    AMBER_TEXT,
    BLUE,
    BLUE_FILL,
    BLUE_TEXT,
    BODY,
    DARK,
    GREEN,
    GREEN_TEXT,
    GREY,
    MID_GREY,
    NAVY,
    TEAL,
    WHITE,
    add_bottom_rule,
    add_callout,
    add_page_field,
    configure_document,
    prevent_row_split,
    set_cell_margins,
    set_cell_shading,
    set_cell_width,
    set_repeat_table_header,
    set_table_borders,
    set_table_width,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Kompliance_Questions_for_Marcelo_2026-07-22.docx"


GATE_STYLE = {
    "Pilot sign-off": (GREEN, GREEN_TEXT),
    "Commercial launch": (AMBER, AMBER_TEXT),
    "Roadmap": (BLUE_FILL, BLUE_TEXT),
}


SECTIONS = [
    (
        "1. Product scope and acceptance",
        "Confirm what Marcelo expects this release to be and the decision he will make after testing.",
        [
            ("1.1", "Is the immediate target a private single-company pilot, a production release for one company, or a multi-company commercial platform?", "The current release is a controlled private pilot.", "Choose one and explain the intended users.", "Pilot sign-off"),
            ("1.2", "Which parts of the platform must work for Marcelo to accept the current pilot?", "Acceptance needs a defined minimum scope.", "List the must-pass workflows.", "Pilot sign-off"),
            ("1.3", "Which features may be accepted with conditions or deferred to a later phase?", "This prevents non-blocking roadmap work from delaying the pilot.", "List acceptable deferrals and target phase.", "Pilot sign-off"),
            ("1.4", "Who is the business owner authorised to accept or reject the pilot?", "A named decision-maker is required.", "Name, role and contact method.", "Pilot sign-off"),
            ("1.5", "Who will be the technical owner for defects, releases and operational decisions?", "The platform needs a named technical counterpart.", "Name, role and contact method.", "Pilot sign-off"),
            ("1.6", "Which company, sites and teams should participate in the first acceptance test?", "Testing should cover representative operations without widening data access unnecessarily.", "Company, site and team list.", "Pilot sign-off"),
            ("1.7", "What date should pilot testing start and when is the acceptance decision expected?", "A defined window is needed for triage and release planning.", "Start date, end date and decision meeting.", "Pilot sign-off"),
            ("1.8", "What measurable outcomes should demonstrate value?", "Marcelo listed reduced repeated entry, faster onboarding and improved compliance visibility.", "Define baseline and target measures.", "Commercial launch"),
        ],
    ),
    (
        "2. Users, roles and account access",
        "Confirm who may use the platform and exactly what each role can see or change.",
        [
            ("2.1", "Which user roles are required at launch?", "The current product includes Administrator, Editor and Viewer roles plus worker access.", "Confirm roles and add any missing roles.", "Pilot sign-off"),
            ("2.2", "What may an Administrator create, edit, delete, export and approve?", "Administrative power must be explicit.", "Describe allowed and prohibited actions.", "Commercial launch"),
            ("2.3", "What may an Editor create or change, and which actions require approval?", "Current Editors can manage local workflows but not users or system settings.", "Confirm or correct the permission boundary.", "Pilot sign-off"),
            ("2.4", "Should Viewers be allowed to download documents, or only view them in the browser?", "Viewing and downloading are separate data-access decisions.", "Choose: view only / view and download / category-based.", "Pilot sign-off"),
            ("2.5", "Which roles may access personal, medical, emergency-contact and certification information?", "Sensitive worker fields need least-privilege rules.", "Role-to-field access matrix.", "Commercial launch"),
            ("2.6", "Is authenticator-app MFA mandatory for Administrators, Editors, all company users, or optional?", "MFA is implemented but enrolment policy is a business decision.", "Choose users, enforcement date and recovery owner.", "Pilot sign-off"),
            ("2.7", "Is mobile-number authentication required, and if so should it use SMS OTP, WhatsApp, or another approved method?", "Email authentication exists; mobile authentication needs a provider and privacy decision.", "Choose method/provider or mark not required.", "Commercial launch"),
        ],
    ),
    (
        "3. Universal worker profile and consent",
        "Validate the worker record, QR journey, company access and revocation rules.",
        [
            ("3.1", "Which worker-profile fields are mandatory, optional, or prohibited?", "Marcelo requested identity, contact, emergency, employment, skills, trade, qualifications, medical, certification, training and documents.", "Mark each field group mandatory/optional/prohibited.", "Pilot sign-off"),
            ("3.2", "Should workers be allowed to edit every profile field themselves?", "Some identity or compliance fields may need company verification.", "List self-editable and verified-only fields.", "Commercial launch"),
            ("3.3", "Should optional medical information be stored at all?", "Medical data is special-category personal data and needs a clear lawful purpose.", "Choose: no / yes; provide purpose, access and retention.", "Commercial launch"),
            ("3.4", "Which profile fields may appear on the public QR page before company access is approved?", "The current design lets the worker choose public fields.", "Confirm default-public, optional-public and never-public fields.", "Pilot sign-off"),
            ("3.5", "Should a company access request expire if the worker does not respond?", "Pending requests need a lifecycle.", "Choose expiry period and reminder rule.", "Commercial launch"),
            ("3.6", "May a worker approve different fields for different companies?", "Field-level, company-specific consent is implemented.", "Confirm yes/no and any mandatory company fields.", "Pilot sign-off"),
            ("3.7", "What must happen immediately when a worker revokes access?", "Current design stops company and API access immediately while preserving audit history.", "Confirm visibility, notification and audit expectations.", "Pilot sign-off"),
            ("3.8", "May a company retain a previously imported copy after consent is revoked?", "This affects data minimisation, operational records and legal obligations.", "Choose policy and lawful basis.", "Commercial launch"),
            ("3.9", "Who verifies worker identity, qualifications and uploaded evidence?", "The platform separates worker-submitted data from company review.", "Name responsible role and verification rules.", "Commercial launch"),
            ("3.10", "Should workers be able to close their account, and what information must be retained afterward?", "Account closure and statutory/audit retention need defined rules.", "Describe deletion, anonymisation and retention.", "Commercial launch"),
        ],
    ),
    (
        "4. Documents, viewing and expiry",
        "Confirm document categories, permissions, expiry rules and handling of protected imports.",
        [
            ("4.1", "Which document categories are required for workers at launch?", "Marcelo listed GA1, GA2, GA3, AF3, RAMS, inductions, certifications, licences and medical certificates.", "Confirm, rename, add or remove categories.", "Pilot sign-off"),
            ("4.2", "Which company document categories are required?", "Marcelo listed plant, equipment and site documents.", "Confirm category and ownership structure.", "Commercial launch"),
            ("4.3", "Which file types and maximum file sizes should be accepted for each category?", "Current workflows validate common PDF, office, CSV and image formats.", "Provide category/type/size rules.", "Pilot sign-off"),
            ("4.4", "Who may replace or delete a writable document, and when should deletion be blocked?", "Imported source files remain read-only; local files can be governed separately.", "Define role, state and retention restrictions.", "Commercial launch"),
            ("4.5", "How should version history be displayed, and which version is considered current?", "Automatic versions exist; the preferred presentation needs confirmation.", "Describe current-version and history rules.", "Roadmap"),
            ("4.6", "Should users be able to view and download every archived PDF, or should restrictions vary by role/category?", "GA1/GA2/GA3 viewing and download are implemented.", "Provide role/category matrix.", "Pilot sign-off"),
            ("4.7", "What exact expiry thresholds and colours should apply?", "Current convention is valid above 30 days, due soon within 30 days and expired after the date.", "Confirm days, colours and missing-date state.", "Pilot sign-off"),
            ("4.8", "Which document categories require automatic expiry extraction?", "Best-effort extraction exists; scanned images may require OCR.", "List categories and acceptable confidence/manual review.", "Commercial launch"),
            ("4.9", "Is OCR for scanned images required before launch?", "OCR/provider work is not required for the current private pilot unless Marcelo makes it a gate.", "Choose: launch gate / later phase / not required.", "Commercial launch"),
            ("4.10", "How long must documents and document versions be retained?", "Retention must match operational and legal requirements.", "Retention by category and deletion authority.", "Commercial launch"),
        ],
    ),
    (
        "5. Compliance, supervisor review and approvals",
        "Confirm review states, routing, decision authority and escalation.",
        [
            ("5.1", "Which requests must the platform support at launch?", "Current types include inspection, renewal, approval, missing document and additional information.", "Confirm request catalogue and priorities.", "Pilot sign-off"),
            ("5.2", "Which departments exist and who is the responsible contact for each?", "Routing requires approved Safety, HR, Plant, Training and Administration contacts or their replacements.", "Department, primary contact and backup contact.", "Pilot sign-off"),
            ("5.3", "How should requests be routed when a department contact is absent or inactive?", "Fallback ownership prevents unhandled requests.", "Define fallback and escalation order.", "Commercial launch"),
            ("5.4", "What is the required supervisor document-review status model?", "Unread/viewed plus approve/decline history are implemented.", "Confirm states, highlighting and when highlighting clears.", "Pilot sign-off"),
            ("5.5", "Which uploaded items require formal approval rather than only a viewed status?", "Different categories may need different controls.", "List category and required decision.", "Commercial launch"),
            ("5.6", "Who may approve or decline an induction?", "The platform records reviewer, date, decision, comments and history.", "Role/site/company approval matrix.", "Pilot sign-off"),
            ("5.7", "Are comments mandatory when declining or requesting additional information?", "Mandatory reasons improve audit quality.", "Choose rule and minimum information.", "Pilot sign-off"),
            ("5.8", "Can an approved induction later be withdrawn, expired or superseded?", "A complete lifecycle needs post-approval rules.", "Define states, authority and notifications.", "Commercial launch"),
            ("5.9", "What service levels should apply to document and induction reviews?", "Escalations and reminders need target response times.", "Hours/days by request type and priority.", "Commercial launch"),
            ("5.10", "Which compliance events must appear in the audit history and for how long?", "Current mutations and decisions are audited.", "Confirm event catalogue and retention period.", "Commercial launch"),
        ],
    ),
    (
        "6. Notifications, email, SMS and scheduling",
        "Approve channels, recipients, reminder timing and automated-delivery controls.",
        [
            ("6.1", "Which notifications must be sent in-app, by email, by SMS or by push?", "In-app and Gmail send capability exist; SMS/push need providers.", "Event-to-channel matrix.", "Commercial launch"),
            ("6.2", "May users choose their own notification channels, or are some messages mandatory?", "Preferences are implemented, but mandatory safety notices may override them.", "Define optional and mandatory channels.", "Commercial launch"),
            ("6.3", "Who should receive expiry reminders for worker, company, plant, equipment and site records?", "Recipients must be approved before automation is enabled.", "Role/contact per record type.", "Pilot sign-off"),
            ("6.4", "Which reminder intervals are required?", "The platform supports 7, 14, 30, 60 and 90-day windows.", "Choose intervals by category.", "Pilot sign-off"),
            ("6.5", "Should reminders repeat after the first message, and when should they stop?", "Deduplication exists; escalation cadence needs a business rule.", "Define repeats, escalation and stop condition.", "Commercial launch"),
            ("6.6", "Is the current Gmail sender acceptable for the pilot, and what sender identity should be used commercially?", "A controlled live send was accepted by Gmail; a business sender identity still needs approval.", "Pilot sender and future sender address/display name.", "Pilot sign-off"),
            ("6.7", "Who authorises enabling the automatic scheduler?", "The scheduler is intentionally disabled until recipients, intervals and governance are approved.", "Name approver and prerequisites.", "Pilot sign-off"),
            ("6.8", "Are SMS and push notifications launch requirements or later enhancements?", "No provider has been selected.", "Choose launch gate / later phase / not required.", "Commercial launch"),
            ("6.9", "Who will monitor failed deliveries and retry or escalate them?", "Delivery history and retries exist but need operational ownership.", "Name role and response target.", "Commercial launch"),
        ],
    ),
    (
        "7. Interface, devices, language and accessibility",
        "Confirm the supported user experience and quality bar.",
        [
            ("7.1", "Which desktop browsers, mobile browsers and tablet devices must be officially supported?", "Primary paths were checked on desktop and a narrow mobile viewport.", "List minimum browser/device versions.", "Pilot sign-off"),
            ("7.2", "Which workflows must be fully usable on a phone?", "Tables, QR access, forms, signatures and document viewing have different mobile priorities.", "List mandatory mobile workflows.", "Pilot sign-off"),
            ("7.3", "Who will perform native-language review for English, Portuguese and Spanish?", "Primary translations exist; rare/error states and native review remain.", "Name reviewer per language and target date.", "Commercial launch"),
            ("7.4", "Is formal WCAG accessibility conformance required, and at what level?", "No independent accessibility audit has yet been approved.", "Choose target standard/level and assessor.", "Commercial launch"),
            ("7.5", "Are there brand, terminology or layout changes Marcelo wants before testing?", "The interface has been modernised but business terminology still needs customer confirmation.", "List required changes and examples.", "Pilot sign-off"),
            ("7.6", "Should dates, times, phone numbers and names follow Irish formatting everywhere?", "Localisation rules should be consistent across languages.", "Confirm locale rules and exceptions.", "Commercial launch"),
        ],
    ),
    (
        "8. Data migration and protected customer information",
        "Authorise source systems, reconciliation and the boundary between protected and writable data.",
        [
            ("8.1", "Does the current 3,597-record snapshot represent the complete authorised pilot dataset?", "The imported snapshot is protected and read-only.", "Confirm total or list known omissions.", "Pilot sign-off"),
            ("8.2", "Which representative workers, sites, forms and documents should be used for reconciliation?", "Acceptance should verify known records, not only totals.", "Provide sample IDs/names through a secure channel.", "Pilot sign-off"),
            ("8.3", "Are all 3,077 source files authorised for this private pilot?", "The archive is mounted read-only and should not contain unauthorised material.", "Confirm authority and any exclusions.", "Pilot sign-off"),
            ("8.4", "Which additional client should be migrated next after the pilot?", "Marcelo mentioned Grandbrind as an example; no additional extraction should occur without written approval.", "Client, owner, priority and proposed date.", "Roadmap"),
            ("8.5", "Who can provide written authorisation and source-schema information for each additional client?", "Per-client authority and mapping are mandatory.", "Name owner and approval method.", "Commercial launch"),
            ("8.6", "What reconciliation tolerance is acceptable for records, relationships and attachments?", "Migration packages produce counts and checksums, but acceptance thresholds need definition.", "Define required match and permitted exceptions.", "Commercial launch"),
            ("8.7", "Should failed or partial migrations roll back completely?", "Transactional isolated imports are implemented.", "Confirm all-or-nothing rule and exception handling.", "Commercial launch"),
            ("8.8", "How long should migration packages, source exports and reconciliation reports be retained?", "These artifacts may contain personal data and need lifecycle controls.", "Retention and secure-destruction rule.", "Commercial launch"),
        ],
    ),
    (
        "9. Privacy, legal, security and governance",
        "Obtain decisions that cannot be completed by engineering alone.",
        [
            ("9.1", "Who is the approved privacy contact shown to users?", "The platform needs a real privacy owner and contact channel.", "Name, role, email/phone and escalation route.", "Pilot sign-off"),
            ("9.2", "What is the lawful basis for processing worker profiles, documents, medical data and cross-company sharing?", "Irish/EU privacy compliance requires a documented basis by purpose.", "Legal/privacy owner response by data purpose.", "Commercial launch"),
            ("9.3", "Which entity is controller and which entities are processors for each workflow?", "Multi-company sharing changes controller/processor responsibilities.", "Provide responsibility model and DPA owner.", "Commercial launch"),
            ("9.4", "What retention period applies to accounts, documents, submissions, certificates, messages, notifications and audit logs?", "Current protected records are never removed by local cleanup.", "Retention schedule by record type.", "Commercial launch"),
            ("9.5", "How should access, correction, restriction, portability and erasure requests be handled?", "Data-subject request operations need owner, verification and response timing.", "Process owner and SLA.", "Commercial launch"),
            ("9.6", "Which subprocessors and hosting locations are approved?", "Email, SMS, push, OCR, monitoring and hosting providers require review.", "Approved providers/regions or approval process.", "Commercial launch"),
            ("9.7", "Who is responsible for incident response and notifying affected organisations or individuals?", "A live service needs a documented response chain.", "Names, contact route and escalation times.", "Commercial launch"),
            ("9.8", "Is an independent penetration test required before one-company production, multi-company production, or both?", "Independent assurance remains outstanding.", "Choose release gate, scope and assessor.", "Commercial launch"),
            ("9.9", "What backup retention and recovery objectives are required?", "Verified backups and restore rehearsal exist; RPO/RTO are not yet business-approved.", "RPO, RTO, backup retention and restore owner.", "Commercial launch"),
            ("9.10", "Who may access production logs, backups and administrative audit data?", "Operational data can contain identifiers and security events.", "Role/access matrix and review frequency.", "Commercial launch"),
        ],
    ),
    (
        "10. API and external integrations",
        "Prioritise partner integrations and define the access contract.",
        [
            ("10.1", "Which external induction or contractor systems should integrate first?", "REST resources and an OpenAPI contract exist, but no partner is yet approved.", "System, owner, purpose and target date.", "Roadmap"),
            ("10.2", "Which worker fields and resources may each integration access?", "API access must follow worker consent and tenant scope.", "Partner-to-resource/field matrix.", "Commercial launch"),
            ("10.3", "Should integrations be read-only, or may approved partners create/update data?", "Current sharing focuses on audited access and import.", "Choose permissions by partner/resource.", "Roadmap"),
            ("10.4", "Are current pagination and rate limits acceptable?", "The default token limit is 120 requests per minute and 100 records per page.", "Confirm limits or provide expected volumes.", "Roadmap"),
            ("10.5", "Is a client SDK required, and for which language/platform?", "An SDK is deferred until a real integration is approved.", "Choose language, partner and priority.", "Roadmap"),
            ("10.6", "What security review is required before issuing a production API token?", "Token issuance, revocation and audit are implemented.", "Approval, rotation, expiry and incident rules.", "Commercial launch"),
        ],
    ),
    (
        "11. Public demo, branding and commercial operation",
        "Separate the private customer pilot from any public sales environment and define ownership.",
        [
            ("11.1", "Does Marcelo want a separate public sales demo?", "The current environment uses authorised customer data and must not be presented publicly.", "Choose yes/no and target date.", "Roadmap"),
            ("11.2", "If a public demo is required, which fictional companies, workers, sites and workflows should it contain?", "A convincing demo needs an approved synthetic scenario.", "Describe demo story and personas.", "Roadmap"),
            ("11.3", "What SAMPLE or DEMO watermark wording and appearance should be used?", "All public-demo files must be fictional and clearly marked.", "Wording, placement, opacity and languages.", "Roadmap"),
            ("11.4", "Who approves that a demo dataset and every demo document contain no real personal/company data?", "Public release needs a named PII review owner.", "Name approver and evidence required.", "Commercial launch"),
            ("11.5", "What final product name, logo, colours, support details and legal footer should be used?", "Brand and contact details appear in generated and public-facing content.", "Provide approved brand pack and wording.", "Pilot sign-off"),
            ("11.6", "Who owns onboarding, training, customer support and incident communications after launch?", "Commercial operation needs named processes, not only software.", "Owner and service hours for each function.", "Commercial launch"),
            ("11.7", "What service levels and support channels will be offered to customers?", "Pricing and contracts depend on support commitments.", "Channels, hours, response and resolution targets.", "Commercial launch"),
            ("11.8", "Which Google OAuth verification and mailbox ownership tasks will Marcelo or the business owner handle?", "The OAuth app works but remains unverified and subject to a warning and 100-user cap.", "Name owner and verification target date.", "Commercial launch"),
        ],
    ),
    (
        "12. Testing, defects and release sign-off",
        "Agree how the release will be tested, corrected and formally approved.",
        [
            ("12.1", "Who will test the company portal, worker portal, supervisor workflow and mobile experience?", "Each path needs a named tester.", "Name tester per area and device/browser.", "Pilot sign-off"),
            ("12.2", "Which test accounts and roles should be used?", "Credentials must be supplied separately and never placed in the questionnaire.", "List account purpose/role only.", "Pilot sign-off"),
            ("12.3", "What severity definitions should be used for blocking, high, normal and cosmetic defects?", "Release decisions need consistent triage.", "Define impact and response for each severity.", "Pilot sign-off"),
            ("12.4", "Which defect severities must be closed before acceptance?", "Accepted-with-conditions needs a clear rule.", "Choose blocking criteria and permitted deferrals.", "Pilot sign-off"),
            ("12.5", "Where should defects, screenshots and decisions be recorded?", "Evidence must be controlled and personal information removed where possible.", "Choose tracker/location and access rules.", "Pilot sign-off"),
            ("12.6", "Who approves enabling automated reminders after the controlled email test?", "Scheduler activation is a separate release action.", "Name approver and acceptance evidence.", "Pilot sign-off"),
            ("12.7", "Who signs the business, technical and pilot-user acceptance record?", "All three approvals are listed in the current checklist.", "Name and role for each signature.", "Pilot sign-off"),
            ("12.8", "If the pilot is rejected, what are the rollback trigger and communication steps?", "Verified rollback points exist; the decision process needs approval.", "Trigger, owner, notification list and decision time.", "Pilot sign-off"),
            ("12.9", "After acceptance, what is the exact next release scope?", "Single-company production, multi-company rollout and public demo have different gates.", "Choose scope, target users and target date.", "Commercial launch"),
        ],
    ),
]


def add_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("KOMPLIANCE  |  QUESTIONS FOR MARCELO")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_bottom_rule(p, MID_GREY, "6")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential — decision questionnaire  •  Page ")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("667085")
    add_page_field(p, "PAGE")
    p.add_run(" of ")
    add_page_field(p, "NUMPAGES")


def add_masthead(doc: Document, question_count: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("KOMPLIANCE PLATFORM")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Questions for Marcelo")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Product Decisions, Approvals & Acceptance")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    add_bottom_rule(p, TEAL, "20")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2100, 7260]
    rows = [
        ("Prepared for", "Marcelo — customer and product review"),
        ("Status date", "22 July 2026"),
        ("Question set", f"{question_count} traceable questions across {len(SECTIONS)} sections"),
        ("Purpose", "Record decisions required for pilot acceptance and commercial release"),
    ]
    for row, values in zip(table.rows, rows):
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, values):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=55, bottom=55)
            cell.text = value
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], GREY)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("667085")
    set_table_borders(table, "E4E7EC", "2")


def add_gate_summary(doc: Document) -> None:
    doc.add_heading("How to answer", level=1)
    add_callout(
        doc,
        "Recommended approach",
        "Marcelo can answer directly, nominate the correct owner, or mark an item Not required. Questions marked Pilot sign-off should be resolved first; commercial and roadmap questions can be scheduled with named owners and dates.",
    )
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2400, 3480, 3480]
    headers = ["Gate", "Meaning", "Expected response"]
    for cell, width, label in zip(table.rows[0].cells, widths, headers):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    entries = [
        ("Pilot sign-off", "Needed to test or accept the current controlled pilot.", "Answer before or during UAT."),
        ("Commercial launch", "Needed before wider production or paid customer operation.", "Decide, assign an owner and target date."),
        ("Roadmap", "Affects later integrations, demo or product expansion.", "Prioritise or mark Not required."),
    ]
    for gate, meaning, response in entries:
        row = table.add_row()
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, (gate, meaning, response)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9)
        fill, color = GATE_STYLE[gate]
        set_cell_shading(row.cells[0], fill)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(color)
    set_table_borders(table)

    doc.add_heading("Response conventions", level=2)
    for text in [
        "Do not place passwords, OAuth values, recovery codes, private keys or other credentials in this document.",
        "Where personal data examples are needed, provide them through the agreed secure channel and reference only the record identifier here.",
        "If Marcelo is not the decision owner, record the owner’s name and a target date instead of leaving the question unanswered.",
        "Use Not required only when the feature or control is explicitly outside the agreed release scope.",
    ]:
        doc.add_paragraph(text, style="List Bullet")


def add_question_table(doc: Document, title: str, intro: str, questions: list[tuple[str, str, str, str, str]]) -> None:
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.keep_with_next = True
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [620, 4360, 3040, 1340]
    headers = ["ID", "Question and reason", "Marcelo response", "Gate"]
    for cell, width, label in zip(table.rows[0].cells, widths, headers):
        set_cell_width(cell, width)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])

    for qid, question, reason, response_prompt, gate in questions:
        row = table.add_row()
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=90, bottom=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        row.cells[0].text = qid
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(8)

        qcell = row.cells[1]
        qcell.text = ""
        p = qcell.paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(question)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(DARK)
        p = qcell.add_paragraph(reason)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.font.size = Pt(7.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor.from_string("667085")

        rcell = row.cells[2]
        rcell.text = ""
        p = rcell.paragraphs[0]
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("Requested response: ")
        r.font.bold = True
        r.font.size = Pt(7.5)
        r = p.add_run(response_prompt)
        r.font.size = Pt(7.5)
        p = rcell.add_paragraph("Answer / owner / target date:")
        p.paragraph_format.space_after = Pt(0)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(7.5)

        row.cells[3].text = gate
        p = row.cells[3].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        fill, color = GATE_STYLE[gate]
        set_cell_shading(row.cells[3], fill)
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(7.5)
        p.runs[0].font.color.rgb = RGBColor.from_string(color)
    set_table_borders(table)


def add_completion_record(doc: Document) -> None:
    doc.add_heading("Questionnaire completion record", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2600, 6760]
    rows = [
        ("Current status", "Awaiting Marcelo’s responses"),
        ("Response owner", "Marcelo or nominated owner for each question"),
        ("Pilot questions due", "Before or during the customer acceptance session"),
        ("Commercial questions due", "Before wider production or commercial release approval"),
        ("Decision record", "Transfer accepted decisions into the release checklist and project issue tracker"),
    ]
    for row, values in zip(table.rows, rows):
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, values):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=100, bottom=100)
            cell.text = value
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], GREY)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_table_borders(table)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("END OF QUESTIONNAIRE")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("667085")


def build() -> Path:
    question_count = sum(len(items) for _, _, items in SECTIONS)
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Kompliance – Questions for Marcelo"
    doc.core_properties.subject = "Product decisions, approvals and acceptance questionnaire"
    doc.core_properties.author = "Kompliance Project Team"
    doc.core_properties.comments = "Companion questionnaire to Marcelo's requirements and delivery status review."
    add_running_header_footer(doc)
    add_masthead(doc, question_count)
    add_gate_summary(doc)
    doc.add_page_break()
    for title, intro, questions in SECTIONS:
        add_question_table(doc, title, intro, questions)
    add_completion_record(doc)
    doc.save(OUTPUT)
    print(f"{OUTPUT}\nquestions={question_count}")
    return OUTPUT


if __name__ == "__main__":
    build()
