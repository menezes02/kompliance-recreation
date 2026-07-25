from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from build_marcelo_questions import SECTIONS, add_question_table
from build_marcelo_review import (
    AMBER,
    AMBER_TEXT,
    BLUE,
    DARK,
    GREY,
    MID_GREY,
    NAVY,
    TEAL,
    add_bottom_rule,
    add_callout,
    add_page_field,
    configure_document,
    prevent_row_split,
    set_cell_margins,
    set_cell_shading,
    set_cell_width,
    set_table_borders,
    set_repeat_table_header,
    set_table_width,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Kompliance_Priority_Answers_Required_from_Marcelo_2026-07-22.docx"


QUESTION_LOOKUP = {
    item[0]: item
    for _, _, questions in SECTIONS
    for item in questions
}


GROUPS = [
    (
        "1. Acceptance scope and decision owners",
        "These answers define what Marcelo is accepting, who may decide, and when the pilot can close.",
        ["1.2", "1.4", "1.5", "1.6", "1.7", "12.3", "12.4", "12.7", "12.9"],
    ),
    (
        "2. Roles and permissions",
        "These answers determine who may view, download, change, approve and administer information.",
        ["2.1", "2.2", "2.3", "2.4", "2.5", "2.6"],
    ),
    (
        "3. Retention, privacy and sensitive information",
        "These answers are required to approve how personal information and records are stored and removed.",
        ["3.3", "3.10", "4.10", "8.8", "9.1", "9.2", "9.3", "9.4", "9.5"],
    ),
    (
        "4. Responsible contacts and routing",
        "These answers ensure requests, reviews, incidents and delivery failures always reach a named owner.",
        ["5.2", "5.3", "5.6", "6.3", "6.9", "9.7"],
    ),
    (
        "5. Notification and scheduler rules",
        "These answers must be approved before automated reminders are enabled.",
        ["6.1", "6.2", "6.4", "6.5", "6.6", "6.7", "6.8", "12.6"],
    ),
]


def add_running_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("KOMPLIANCE  |  PRIORITY ANSWERS FROM MARCELO")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_bottom_rule(p, MID_GREY, "6")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential — priority decision form  •  Page ")
    r.font.name = "Calibri"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("667085")
    add_page_field(p, "PAGE")
    p.add_run(" of ")
    add_page_field(p, "NUMPAGES")


def add_masthead(doc: Document, count: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("KOMPLIANCE PLATFORM")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Priority Answers Required")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Marcelo Decision Form")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    add_bottom_rule(p, TEAL, "20")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2100, 7260]
    values = [
        ("Prepared for", "Marcelo — priority customer decisions"),
        ("Status date", "22 July 2026"),
        ("Decision set", f"{count} priority answers across five blocking areas"),
        ("Immediate outcome", "Confirm pilot ownership and remove avoidable release uncertainty"),
    ]
    for row, pair in zip(table.rows, values):
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, pair):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=55, bottom=55)
            cell.text = value
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], GREY)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("667085")
    set_table_borders(table, "E4E7EC", "2")
    set_repeat_table_header(table.rows[0])


def add_intro(doc: Document) -> None:
    doc.add_heading("Why these answers are needed", level=1)
    add_callout(
        doc,
        "Focused decision pack",
        "This document extracts only the priority decisions from the full 101-question questionnaire. Marcelo can answer directly, nominate the correct owner, or give a target date. These answers do not replace legal or security review where specialist approval is required.",
    )
    set_repeat_table_header(doc.tables[-1].rows[0])
    doc.add_paragraph()
    for item in [
        "Answer pilot sign-off items first; they determine the acceptance session and immediate release decision.",
        "For commercial-launch items, record the responsible owner and target date if the final decision is not yet available.",
        "Do not enter passwords, OAuth values, recovery codes, private keys or personal-data examples in this document.",
        "Accepted answers should be copied into the release checklist, operating procedures and project issue tracker.",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def add_completion_record(doc: Document) -> None:
    doc.add_heading("Decision completion record", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    widths = [2700, 6660]
    rows = [
        ("Current status", "Awaiting Marcelo’s priority answers"),
        ("Pilot decision owner", "To be named in section 1"),
        ("Technical decision owner", "To be named in section 1"),
        ("Notification approval", "Scheduler remains disabled until section 5 is approved"),
        ("Next action", "Review responses, update the acceptance checklist and assign unresolved commercial items"),
    ]
    for row, values in zip(table.rows, rows):
        prevent_row_split(row)
        for cell, width, value in zip(row.cells, widths, values):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=100, bottom=100)
            cell.text = value
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], GREY)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    add_callout(
        doc,
        "Release control",
        "The platform remains suitable for controlled testing. Commercial launch and automated reminders remain conditional on the approvals recorded in this form and the release checklist.",
        fill=AMBER,
        accent=AMBER_TEXT,
    )
    set_repeat_table_header(doc.tables[-1].rows[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("END OF PRIORITY DECISION FORM")
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("667085")


def build() -> Path:
    count = sum(len(ids) for _, _, ids in GROUPS)
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Kompliance – Priority Answers Required from Marcelo"
    doc.core_properties.subject = "Focused roles, permissions, retention, contacts and notification decision form"
    doc.core_properties.author = "Kompliance Project Team"
    doc.core_properties.comments = "Focused companion to the full Questions for Marcelo questionnaire."
    add_running_header_footer(doc)
    add_masthead(doc, count)
    add_intro(doc)
    doc.add_page_break()
    for title, intro, ids in GROUPS:
        questions = [QUESTION_LOOKUP[qid] for qid in ids]
        add_question_table(doc, title, intro, questions)
    add_completion_record(doc)
    doc.save(OUTPUT)
    print(f"{OUTPUT}\npriority_questions={count}")
    return OUTPUT


if __name__ == "__main__":
    build()
